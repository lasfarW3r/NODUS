from pathlib import Path

from docx import Document as DocxDocument

from backend.models.document import Document


def parse_docx(file: Path) -> Document:
    doc = DocxDocument(file)

    content = ""

    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    return Document(
        path=file,
        extension=file.suffix,
        size=file.stat().st_size,
        modified_time=file.stat().st_mtime,
        content=content,
    )